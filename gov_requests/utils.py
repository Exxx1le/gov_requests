import os
from datetime import timedelta
from pathlib import Path

import docx
import requests
from docx import Document  # Для создания DOC-файла

from constants import AUTHORITY_ID, AUTHORITY_NAMES

# URL для запроса документов
API_URL = "http://publication.pravo.gov.ru/api/Documents"


def fetch_documents_by_date(date, authority_id):
    """
    Fetches documents for a specified date and government authority.
    Excludes resolutions (documentTypeId=7ff5b3b5-3757-44f1-bb76-3766cabe3593)
    for all authorities.

    Args:
        date (str): Date in DD.MM.YYYY format
        authority_id (str): GUID of the government authority

    Returns:
        list: List of found documents (excluding resolutions)
    """
    params = {
        "periodType": "day",
        "date": date,
        "SignatoryAuthorityId": authority_id
    }

    try:
        response = requests.get(API_URL, params=params)
        response.raise_for_status()
        data = response.json()

        # Фильтруем результаты, исключая распоряжения
        return [
            doc for doc in data.get("items", [])
            if doc.get("documentTypeId") != "7ff5b3b5-3757-44f1-bb76-3766cabe3593"
        ]
    except requests.exceptions.RequestException as e:
        print(f"Ошибка при выполнении запроса для даты {date}: {e}")
        return []


def add_hyperlink(paragraph, text, url):
    """
    Adds a clickable hyperlink to a paragraph in a Word document.

    Args:
        paragraph: Document paragraph object
        text (str): Display text for the hyperlink
        url (str): URL for the hyperlink

    Returns:
        OxmlElement: Created hyperlink element
    """
    # Получаем связь с word документом
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Создаем гиперссылку
    hyperlink = docx.oxml.shared.OxmlElement("w:hyperlink")
    hyperlink.set(docx.oxml.shared.qn("r:id"), r_id)

    # Создаем элемент run
    new_run = docx.oxml.shared.OxmlElement("w:r")
    rPr = docx.oxml.shared.OxmlElement("w:rPr")

    # Добавляем стиль для гиперссылки
    rStyle = docx.oxml.shared.OxmlElement("w:rStyle")
    rStyle.set(docx.oxml.shared.qn("w:val"), "Hyperlink")
    rPr.append(rStyle)

    # Добавляем текст
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)
    return hyperlink


def save_to_doc(documents_by_authority, filename):
    """
    Saves documents to a .docx file, grouping them by government authority.
    Only authorities with documents will be included in the output file.

    Args:
        documents_by_authority (dict): Dictionary where keys are authority IDs and values are lists of documents
        filename (str): Name of the output file

    Returns:
        None
    """
    home_dir = str(Path.home())
    output_path = os.path.join(home_dir, filename)

    doc = Document()

    for authority_id, documents in documents_by_authority.items():
        # Используем словарь для получения названия органа власти
        authority_name = AUTHORITY_NAMES.get(authority_id, authority_id)
        doc.add_heading(authority_name, level=1)

        for doc_info in documents:
            eo_number = doc_info.get("eoNumber", "N/A")
            document_name = doc_info.get("complexName", "Название не указано")
            document_date = doc_info.get("documentDate", "Дата не указана")

            # Добавляем название документа
            doc.add_paragraph(f"{document_name} ({document_date})")

            # Создаем параграф для ссылки
            p = doc.add_paragraph()
            link = f"http://publication.pravo.gov.ru/document/{eo_number}"
            add_hyperlink(p, f"Ссылка на документ", link)

            # Добавляем пустую строку между документами
            doc.add_paragraph("")

        # Добавляем разделитель между органами власти
        doc.add_page_break()

    doc.save(output_path)


def generate_date_range(start_date, end_date):
    """
    Generates a list of dates between start_date and end_date (inclusive).

    Args:
        start_date (datetime.date): Start date
        end_date (datetime.date): End date

    Returns:
        list: List of dates in DD.MM.YYYY format
    """
    dates = []
    current_date = start_date
    while current_date <= end_date:
        dates.append(current_date.strftime("%d.%m.%Y"))
        current_date += timedelta(days=1)
    return dates


def process_documents(start_date, end_date):
    """
    Main function for fetching and processing documents.
    Creates a Word document with documents grouped by authority.

    Args:
        start_date (datetime.date): Start date for document search
        end_date (datetime.date): End date for document search

    Returns:
        None
    """
    current_date = start_date
    all_documents_by_authority = {}

    while current_date <= end_date:
        formatted_date = current_date.strftime("%d.%m.%Y")

        for authority_id in AUTHORITY_ID:
            documents = fetch_documents_by_date(formatted_date, authority_id)
            if documents:
                if authority_id not in all_documents_by_authority:
                    all_documents_by_authority[authority_id] = []
                all_documents_by_authority[authority_id].extend(documents)

        current_date += timedelta(days=1)

    save_to_doc(all_documents_by_authority, 
                f"Документы за {start_date.strftime('%Y-%m-%d')} - {end_date.strftime('%Y-%m-%d')}.docx")
